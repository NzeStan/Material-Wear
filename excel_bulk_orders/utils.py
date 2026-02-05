# excel_bulk_orders/utils.py
"""
Utility functions for Excel Bulk Orders.

Handles:
- Coupon code generation
- Excel template generation with data validation
- Excel file parsing and validation
- Document generation (PDF, Word, Excel) for participants
"""
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.worksheet.datavalidation import DataValidation
from io import BytesIO
import pandas as pd
import string
import random
import logging
from django.conf import settings
from .models import ExcelBulkOrder, ExcelParticipant

logger = logging.getLogger(__name__)


# ============================================================================
# COUPON GENERATION
# ============================================================================

def generate_excel_coupon_codes(bulk_order, count=10):
    """
    Generate unique coupon codes for an Excel bulk order.
    
    Args:
        bulk_order: ExcelBulkOrder instance
        count: Number of coupons to generate (default 10)
    
    Returns:
        List of ExcelCouponCode instances
    
    Raises:
        Exception: If coupon generation fails
    """
    from .models import ExcelCouponCode
    
    chars = string.ascii_uppercase + string.digits
    codes = []
    try:
        for _ in range(count):
            while True:
                code = "".join(random.choices(chars, k=8))
                if not ExcelCouponCode.objects.filter(code=code).exists():
                    coupon = ExcelCouponCode.objects.create(bulk_order=bulk_order, code=code)
                    codes.append(coupon)
                    break
        logger.info(f"Generated {count} coupon codes for Excel bulk order: {bulk_order.reference}")
        return codes
    except Exception as e:
        logger.error(f"Error generating coupon codes: {str(e)}")
        raise


# ============================================================================
# EXCEL TEMPLATE GENERATION
# ============================================================================

def generate_excel_template(bulk_order):
    """
    Generate Excel template with data validation and formatting.
    
    Args:
        bulk_order: ExcelBulkOrder instance
    
    Returns:
        BytesIO: Excel file buffer
    """
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Participants"
    
    # Define styles
    header_font = Font(bold=True, size=12, color="FFFFFF")
    header_fill = PatternFill(start_color="064E3B", end_color="064E3B", fill_type="solid")
    header_alignment = Alignment(horizontal="center", vertical="center")
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # Column headers
    headers = ['S/N', 'Full Name', 'Size']
    if bulk_order.requires_custom_name:
        headers.append('Custom Name')
    headers.append('Coupon Code')
    
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num)
        cell.value = header
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment
        cell.border = border
    
    # Size validation dropdown
    size_options = '"Small,Medium,Large,Extra Large,2X Large,3X Large,4X Large"'
    size_col = chr(ord('A') + 2)  # Column C
    dv = DataValidation(type="list", formula1=size_options, allow_blank=False)
    dv.error = 'Please select a size from the dropdown'
    dv.errorTitle = 'Invalid Size'
    ws.add_data_validation(dv)
    dv.add(f'{size_col}2:{size_col}1000')
    
    # Example row
    example_data = [1, 'John Doe', 'Medium']
    if bulk_order.requires_custom_name:
        example_data.append('JOHN')
    example_data.append('')  # Coupon Code
    
    for col_num, value in enumerate(example_data, 1):
        cell = ws.cell(row=2, column=col_num)
        cell.value = value
        cell.border = border
        cell.alignment = Alignment(horizontal="left", vertical="center")
    
    # Instructions worksheet
    instructions_ws = wb.create_sheet("Instructions")
    instructions = [
        f"Excel Bulk Order Template - {bulk_order.title}",
        "",
        "HOW TO FILL THIS TEMPLATE:",
        "",
        "1. S/N: Auto-numbered (1, 2, 3, ...)",
        "   - Use sequential numbers starting from 1",
        "",
        "2. Full Name: Participant's complete name",
        "   - Required field",
        "   - Example: 'John Doe', 'Mary Jane Smith'",
        "",
        "3. Size: Uniform/shirt size",
        "   - Select from dropdown only",
        "   - Options: Small, Medium, Large, Extra Large, 2X Large, 3X Large, 4X Large",
        "",
    ]
    
    if bulk_order.requires_custom_name:
        instructions.extend([
            "4. Custom Name: Name for badge/tag",
            "   - Required field",
            "   - Keep it short (max 20 characters recommended)",
            "   - Example: 'JD', 'Mary', 'Coach John'",
            "",
            "5. Coupon Code: Optional discount code",
        ])
    else:
        instructions.extend([
            "4. Coupon Code: Optional discount code",
        ])
    
    instructions.extend([
        "   - Optional field",
        "   - If valid, participant is FREE",
        "   - Leave blank if no coupon",
        "",
        "IMPORTANT NOTES:",
        "- Do NOT delete or rename column headers",
        "- Do NOT skip rows",
        "- Do NOT add extra columns",
        "- Delete the example row before filling your data",
        "- Save as .xlsx format",
        "",
        f"After filling, upload the file to complete your order.",
        "",
        f"Questions? Contact: {settings.COMPANY_EMAIL}",
    ])
    
    for row_num, instruction in enumerate(instructions, 1):
        cell = instructions_ws.cell(row=row_num, column=1)
        cell.value = instruction
        if row_num == 1:
            cell.font = Font(bold=True, size=14)
        elif instruction.startswith(("1.", "2.", "3.", "4.", "5.")):
            cell.font = Font(bold=True)
    
    # Adjust column widths
    ws.column_dimensions['A'].width = 8  # S/N
    ws.column_dimensions['B'].width = 30  # Full Name
    ws.column_dimensions['C'].width = 15  # Size
    if bulk_order.requires_custom_name:
        ws.column_dimensions['D'].width = 20  # Custom Name
        ws.column_dimensions['E'].width = 15  # Coupon Code
    else:
        ws.column_dimensions['D'].width = 15  # Coupon Code
    
    instructions_ws.column_dimensions['A'].width = 70
    
    # Save to BytesIO
    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    logger.info(f"Generated Excel template for bulk order: {bulk_order.reference}")
    return buffer


# ============================================================================
# EXCEL VALIDATION
# ============================================================================

def validate_excel_file(bulk_order, excel_file):
    """
    Validate uploaded Excel file.
    
    Args:
        bulk_order: ExcelBulkOrder instance
        excel_file: Uploaded file object
    
    Returns:
        dict: {
            'valid': bool,
            'errors': list of error dicts,
            'summary': {total_rows, valid_rows, error_rows}
        }
    """
    errors = []
    valid_rows = []
    
    try:
        # Read Excel file
        df = pd.read_excel(excel_file, sheet_name='Participants')
        
        # Expected columns
        expected_columns = ['S/N', 'Full Name', 'Size']
        if bulk_order.requires_custom_name:
            expected_columns.append('Custom Name')
        expected_columns.append('Coupon Code')
        
        # Check column structure
        if list(df.columns) != expected_columns:
            return {
                'valid': False,
                'errors': [{
                    'row': 0,
                    'field': 'Structure',
                    'error': f'Invalid column structure. Expected: {", ".join(expected_columns)}',
                    'current_value': ''
                }],
                'summary': {'total_rows': 0, 'valid_rows': 0, 'error_rows': 1}
            }
        
        # Valid size options
        valid_sizes = {'Small', 'Medium', 'Large', 'Extra Large', '2X Large', '3X Large', '4X Large'}
        
        # Validate each row
        for idx, row in df.iterrows():
            row_num = idx + 2  # Excel rows start at 1, header is row 1
            row_errors = []
            
            # Validate Full Name
            if pd.isna(row['Full Name']) or not str(row['Full Name']).strip():
                row_errors.append({
                    'row': row_num,
                    'field': 'Full Name',
                    'error': 'Full Name is required',
                    'current_value': ''
                })
            
            # Validate Size
            if pd.isna(row['Size']) or str(row['Size']).strip() not in valid_sizes:
                row_errors.append({
                    'row': row_num,
                    'field': 'Size',
                    'error': f'Size must be one of: {", ".join(valid_sizes)}',
                    'current_value': str(row['Size']) if not pd.isna(row['Size']) else ''
                })
            
            
            # Collect errors or mark as valid
            if row_errors:
                errors.extend(row_errors)
            else:
                valid_rows.append(row_num)
        
        total_rows = len(df)
        is_valid = len(errors) == 0 and total_rows > 0
        
        result = {
            'valid': is_valid,
            'errors': errors,
            'summary': {
                'total_rows': total_rows,
                'valid_rows': len(valid_rows),
                'error_rows': len(set(err['row'] for err in errors))
            }
        }
        
        logger.info(
            f"Validated Excel for {bulk_order.reference}: "
            f"Valid={is_valid}, Total={total_rows}, Errors={len(errors)}"
        )
        
        return result
        
    except Exception as e:
        logger.error(f"Error validating Excel file: {str(e)}")
        return {
            'valid': False,
            'errors': [{
                'row': 0,
                'field': 'File',
                'error': f'Error reading Excel file: {str(e)}',
                'current_value': ''
            }],
            'summary': {'total_rows': 0, 'valid_rows': 0, 'error_rows': 1}
        }


# ============================================================================
# PARTICIPANT CREATION FROM EXCEL
# ============================================================================

def create_participants_from_excel(bulk_order, excel_file):
    """
    Create ExcelParticipant records from validated Excel file.
    Should only be called after successful payment.
    
    Args:
        bulk_order: ExcelBulkOrder instance
        excel_file: Uploaded/validated Excel file
    
    Returns:
        int: Number of participants created
    """
    try:
        df = pd.read_excel(excel_file, sheet_name='Participants')
        
        # Size mapping
        size_map = {
            'Small': 'S',
            'Medium': 'M',
            'Large': 'L',
            'Extra Large': 'XL',
            '2X Large': 'XXL',
            '3X Large': 'XXXL',
            '4X Large': 'XXXXL'
        }
        
        participants_created = 0
        
        for idx, row in df.iterrows():
            row_num = idx + 2
            
            # Extract data
            full_name = str(row['Full Name']).strip()
            size_display = str(row['Size']).strip()
            size_code = size_map.get(size_display, 'M')  # Default to Medium if unknown
            
            # Handle custom name - make uppercase if present
            custom_name = None
            if bulk_order.requires_custom_name:
                custom_name = str(row['Custom Name']).strip().upper() if not pd.isna(row['Custom Name']) else ''
            
            coupon_code = str(row['Coupon Code']).strip() if not pd.isna(row['Coupon Code']) else ''
            
            # Handle coupon
            coupon = None
            is_coupon_applied = False
            if coupon_code:
                try:
                    from .models import ExcelCouponCode
                    coupon = ExcelCouponCode.objects.get(
                        code=coupon_code,
                        bulk_order=bulk_order,
                        is_used=False
                    )
                    is_coupon_applied = True
                    # Mark coupon as used
                    coupon.is_used = True
                    coupon.save()
                except ExcelCouponCode.DoesNotExist:
                    pass  # Coupon not found or already used
            
            # Create participant
            ExcelParticipant.objects.create(
                bulk_order=bulk_order,
                full_name=full_name,
                size=size_code,
                custom_name=custom_name if custom_name else None,
                coupon_code=coupon_code if coupon_code else None,
                coupon=coupon,
                is_coupon_applied=is_coupon_applied,
                row_number=row_num
            )
            
            participants_created += 1
        
        logger.info(
            f"Created {participants_created} participants for bulk order: {bulk_order.reference}"
        )
        
        return participants_created
        
    except Exception as e:
        logger.error(f"Error creating participants from Excel: {str(e)}")
        raise


# ============================================================================
# DOCUMENT GENERATION (PDF, WORD, EXCEL)
# ============================================================================

def generate_participants_pdf(bulk_order, request=None):
    """
    Generate PDF list of all participants.
    
    Args:
        bulk_order: ExcelBulkOrder instance
        request: Optional request object for building absolute URIs
    
    Returns:
        BytesIO: PDF file buffer
    
    Raises:
        ImportError: If WeasyPrint is not installed
        Exception: For other PDF generation errors
    """
    try:
        from weasyprint import HTML
        from django.template.loader import render_to_string
        from django.db.models import Count
        from django.utils import timezone
        
        participants = bulk_order.participants.all().order_by('size', 'full_name')
        size_summary = participants.values('size').annotate(count=Count('size')).order_by('size')
        
        context = {
            'bulk_order': bulk_order,
            'participants': participants,
            'size_summary': size_summary,
            'total_participants': participants.count(),
            'total_paid': participants.filter(is_coupon_applied=False).count(),
            'total_free': participants.filter(is_coupon_applied=True).count(),
            'company_name': settings.COMPANY_NAME,
            'company_address': settings.COMPANY_ADDRESS,
            'company_phone': settings.COMPANY_PHONE,
            'company_email': settings.COMPANY_EMAIL,
            'now': timezone.now(),
        }
        
        html_string = render_to_string('excel_bulk_orders/pdf_template.html', context)
        
        if request:
            html = HTML(string=html_string, base_url=request.build_absolute_uri())
        else:
            html = HTML(string=html_string)
        
        pdf_bytes = html.write_pdf()
        
        buffer = BytesIO(pdf_bytes)
        buffer.seek(0)
        
        logger.info(f"Generated PDF for Excel bulk order: {bulk_order.reference}")
        return buffer
        
    except ImportError as e:
        logger.error(f"WeasyPrint not available: {str(e)}")
        raise ImportError("PDF generation not available. Install GTK+ libraries for WeasyPrint.")
    except Exception as e:
        logger.error(f"Error generating PDF for {bulk_order.reference}: {str(e)}")
        raise


def generate_participants_word(bulk_order):
    """
    Generate Word document list of all participants.
    
    Args:
        bulk_order: ExcelBulkOrder instance
    
    Returns:
        BytesIO: Word document buffer
    
    Raises:
        Exception: If document generation fails
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from django.db.models import Count
        from django.utils import timezone
        
        doc = Document()
        
        # Title
        title = doc.add_heading(bulk_order.title, 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Metadata
        doc.add_paragraph(f"Reference: {bulk_order.reference}")
        doc.add_paragraph(f"Coordinator: {bulk_order.coordinator_name}")
        doc.add_paragraph(f"Generated: {timezone.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph("")
        
        participants = bulk_order.participants.all().order_by('size', 'full_name')
        
        # Summary
        doc.add_heading("Summary", level=1)
        summary_para = doc.add_paragraph()
        summary_para.add_run(f"Total Participants: {participants.count()}\n").bold = True
        summary_para.add_run(f"Paid: {participants.filter(is_coupon_applied=False).count()}\n")
        summary_para.add_run(f"Free (Coupons): {participants.filter(is_coupon_applied=True).count()}\n")
        summary_para.add_run(f"Total Amount: â‚¦{bulk_order.total_amount:,.2f}\n")
        
        doc.add_paragraph("")
        
        # Size breakdown
        doc.add_heading("Size Breakdown", level=1)
        size_summary = participants.values('size').annotate(count=Count('size')).order_by('size')
        
        for size_data in size_summary:
            size = size_data['size']
            count = size_data['count']
            doc.add_paragraph(f"{size}: {count} participant(s)", style='List Bullet')
        
        doc.add_paragraph("")
        
        # Participant list by size
        doc.add_heading("Participant List", level=1)
        
        for size_data in size_summary:
            size = size_data['size']
            size_participants = participants.filter(size=size)
            
            if size_participants.exists():
                doc.add_heading(f"Size: {size} ({size_participants.count()})", level=2)
                
                # Determine columns based on custom name requirement
                if bulk_order.requires_custom_name:
                    table = doc.add_table(rows=1, cols=4)
                    table.style = 'Light Grid Accent 1'
                    
                    # Header row
                    header_cells = table.rows[0].cells
                    header_cells[0].text = 'S/N'
                    header_cells[1].text = 'Full Name'
                    header_cells[2].text = 'Custom Name'
                    header_cells[3].text = 'Status'
                    
                    for idx, participant in enumerate(size_participants, 1):
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(idx)
                        row_cells[1].text = participant.full_name
                        # Only show custom name if it exists
                        row_cells[2].text = participant.custom_name.upper() if participant.custom_name else '-'
                        row_cells[3].text = 'Free (Coupon)' if participant.is_coupon_applied else 'Paid'
                else:
                    table = doc.add_table(rows=1, cols=3)
                    table.style = 'Light Grid Accent 1'
                    
                    # Header row
                    header_cells = table.rows[0].cells
                    header_cells[0].text = 'S/N'
                    header_cells[1].text = 'Full Name'
                    header_cells[2].text = 'Status'
                    
                    for idx, participant in enumerate(size_participants, 1):
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(idx)
                        row_cells[1].text = participant.full_name
                        row_cells[2].text = 'Free (Coupon)' if participant.is_coupon_applied else 'Paid'
                
                doc.add_paragraph()
        
        # ====== NEW SECTION: Custom Names by Size (only if custom names enabled) ======
        if bulk_order.requires_custom_name:
            # Add page break to start on new page
            doc.add_page_break()
            
            doc.add_heading('Custom Names by Size', level=1)
            doc.add_paragraph('This section shows all custom names grouped by size for easy copying.')
            doc.add_paragraph()
            
            # Group participants by size
            for size_data in size_summary:
                size = size_data['size']
                size_participants = participants.filter(size=size).order_by('full_name')
                
                # Get all custom names for this size
                custom_names = [
                    participant.custom_name.upper() if participant.custom_name else participant.full_name.upper()
                    for participant in size_participants
                ]
                
                if custom_names:
                    # Add size header
                    doc.add_heading(f'SIZE: {size}', level=2)
                    
                    # Create table with 5 columns for grid layout
                    num_cols = 5
                    num_rows = (len(custom_names) + num_cols - 1) // num_cols  # Ceiling division
                    
                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    table.style = 'Light Grid Accent 1'
                    
                    # Fill table with custom names
                    for idx, custom_name in enumerate(custom_names):
                        row_idx = idx // num_cols
                        col_idx = idx % num_cols
                        table.rows[row_idx].cells[col_idx].text = custom_name
                    
                    doc.add_paragraph()
        
        # Footer
        doc.add_paragraph("")
        footer_para = doc.add_paragraph()
        footer_para.add_run(f"{settings.COMPANY_NAME}\n").bold = True
        footer_para.add_run(f"{settings.COMPANY_ADDRESS}\n")
        footer_para.add_run(f"ðŸ“ž {settings.COMPANY_PHONE} | ðŸ“§ {settings.COMPANY_EMAIL}")
        
        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        logger.info(f"Generated Word document for Excel bulk order: {bulk_order.reference}")
        return buffer
        
    except Exception as e:
        logger.error(f"Error generating Word document for {bulk_order.reference}: {str(e)}")
        raise


def generate_participants_excel(bulk_order):
    """
    Generate Excel summary of all participants.
    
    Args:
        bulk_order: ExcelBulkOrder instance
    
    Returns:
        BytesIO: Excel file buffer
    
    Raises:
        Exception: If Excel generation fails
    """
    try:
        import xlsxwriter
        from django.db.models import Count
        from django.utils import timezone
        
        buffer = BytesIO()
        workbook = xlsxwriter.Workbook(buffer, {'constant_memory': True})
        
        # Define formats
        header_format = workbook.add_format({
            'bold': True,
            'font_size': 12,
            'bg_color': '#064E3B',
            'font_color': 'white',
            'align': 'center',
            'valign': 'vcenter',
            'border': 1
        })
        
        paid_format = workbook.add_format({
            'bg_color': '#ECFDF5',
            'border': 1,
            'align': 'center',
            'valign': 'vcenter'
        })
        
        free_format = workbook.add_format({
            'bg_color': '#FEF3C7',
            'border': 1,
            'align': 'center',
            'valign': 'vcenter'
        })
        
        normal_format = workbook.add_format({
            'border': 1,
            'align': 'center',
            'valign': 'vcenter'
        })
        
        # Create worksheet
        participants_sheet = workbook.add_worksheet('Participants')
        participants_sheet.set_column('A:A', 8)  # S/N
        participants_sheet.set_column('B:B', 25) # Full Name
        participants_sheet.set_column('C:C', 10) # Size
        
        # Conditionally add Custom Name column
        if bulk_order.requires_custom_name:
            participants_sheet.set_column('D:D', 20) # Custom Name
            participants_sheet.set_column('E:E', 15) # Status
            participants_sheet.set_column('F:F', 15) # Coupon Code
            headers = ['S/N', 'Full Name', 'Size', 'Custom Name', 'Status', 'Coupon Code']
        else:
            participants_sheet.set_column('D:D', 15) # Status
            participants_sheet.set_column('E:E', 15) # Coupon Code
            headers = ['S/N', 'Full Name', 'Size', 'Status', 'Coupon Code']
        
        # Write headers
        for col, header in enumerate(headers):
            participants_sheet.write(0, col, header, header_format)
        
        # Write participant data
        participants = bulk_order.participants.all().order_by('row_number')
        for idx, participant in enumerate(participants, 1):
            row_format = free_format if participant.is_coupon_applied else paid_format
            
            col = 0
            participants_sheet.write(idx, col, idx, normal_format)
            col += 1
            participants_sheet.write(idx, col, participant.full_name, row_format)
            col += 1
            participants_sheet.write(idx, col, participant.size, row_format)
            col += 1
            
            # Only include Custom Name column if required
            if bulk_order.requires_custom_name:
                custom_name_display = participant.custom_name.upper() if participant.custom_name else '-'
                participants_sheet.write(idx, col, custom_name_display, row_format)
                col += 1
            
            status = 'Free (Coupon)' if participant.is_coupon_applied else 'Paid'
            participants_sheet.write(idx, col, status, row_format)
            col += 1
            participants_sheet.write(idx, col, participant.coupon_code or '-', row_format)
        
        workbook.close()
        buffer.seek(0)
        
        logger.info(f"Generated Excel for Excel bulk order: {bulk_order.reference}")
        return buffer
        
    except Exception as e:
        logger.error(f"Error generating Excel for {bulk_order.reference}: {str(e)}")
        raise