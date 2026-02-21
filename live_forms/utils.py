# live_forms/utils.py
"""
Utility functions for Live Forms app.

Mirrors bulk_orders/utils.py A-Z:
  - _get_live_form_with_entries    ≡  _get_bulk_order_with_orders
  - generate_live_form_pdf         ≡  generate_bulk_order_pdf
  - generate_live_form_word        ≡  generate_bulk_order_word
  - generate_live_form_excel       ≡  generate_bulk_order_excel

custom_name columns appear ONLY when custom_branding_enabled=True,
both in document headers and in every data row.
"""
from django.http import HttpResponse
from django.template.loader import render_to_string
from django.conf import settings
from django.utils import timezone
from django.db.models import Count
import logging

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Internal helper  (≡ _get_bulk_order_with_orders)
# ---------------------------------------------------------------------------

def _get_live_form_with_entries(live_form):
    """
    Accepts a LiveFormLink instance or a slug string.
    Returns a LiveFormLink with entries pre-fetched.
    """
    from .models import LiveFormLink

    if isinstance(live_form, str):
        live_form = LiveFormLink.objects.prefetch_related("entries").get(slug=live_form)
    return live_form


# ---------------------------------------------------------------------------
# PDF  (≡ generate_bulk_order_pdf)
# ---------------------------------------------------------------------------

def generate_live_form_pdf(live_form, request=None):
    """
    Generate PDF report for a live form.
    Used by both admin interface and API endpoints.

    Args:
        live_form: LiveFormLink instance or slug string
        request:   Optional request object for building absolute URIs

    Returns:
        HttpResponse with PDF content

    Raises:
        ImportError: If WeasyPrint / GTK+ is not installed
        Exception:   For other generation errors
    """
    try:
        from weasyprint import HTML

        live_form = _get_live_form_with_entries(live_form)
        entries = live_form.entries.all().order_by("serial_number")
        size_summary = (
            entries.values("size").annotate(count=Count("size")).order_by("size")
        )

        context = {
            "live_form": live_form,
            "entries": entries,
            "size_summary": size_summary,
            "total_entries": entries.count(),
            "custom_branding_enabled": live_form.custom_branding_enabled,
            "company_name": settings.COMPANY_NAME,
            "company_address": settings.COMPANY_ADDRESS,
            "company_phone": settings.COMPANY_PHONE,
            "company_email": settings.COMPANY_EMAIL,
            "now": timezone.now(),
        }

        html_string = render_to_string("live_forms/pdf_template.html", context)

        if request:
            html = HTML(string=html_string, base_url=request.build_absolute_uri())
        else:
            html = HTML(string=html_string)

        pdf = html.write_pdf()

        response = HttpResponse(pdf, content_type="application/pdf")
        filename = (
            f"live_form_{live_form.slug}_{timezone.now().strftime('%Y%m%d')}.pdf"
        )
        response["Content-Disposition"] = f'attachment; filename="{filename}"'

        logger.info(f"Generated PDF for live form: {live_form.slug}")
        return response

    except ImportError as e:
        logger.error(f"WeasyPrint not available: {str(e)}")
        raise ImportError(
            "PDF generation not available. Install GTK+ libraries for WeasyPrint."
        )
    except Exception as e:
        logger.error(f"Error generating PDF for live form {live_form}: {str(e)}")
        raise


# ---------------------------------------------------------------------------
# Word  (≡ generate_bulk_order_word)
# ---------------------------------------------------------------------------

def generate_live_form_word(live_form):
    """
    Generate Word document (.docx) for a live form.
    Used by both admin interface and API endpoints.

    Args:
        live_form: LiveFormLink instance or slug string

    Returns:
        HttpResponse with Word document content

    Raises:
        Exception: For document generation errors
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from io import BytesIO

        live_form = _get_live_form_with_entries(live_form)
        entries = live_form.entries.all().order_by("serial_number")

        doc = Document()

        # ── Header ──────────────────────────────────────────────────────
        doc.add_heading(settings.COMPANY_NAME, 0)
        doc.add_heading(f"Live Form: {live_form.organization_name}", level=1)
        doc.add_paragraph(
            f"Generated: {timezone.now().strftime('%B %d, %Y — %I:%M %p')}"
        )
        doc.add_paragraph(
            f"Expires At: {live_form.expires_at.strftime('%B %d, %Y — %I:%M %p')}"
        )
        doc.add_paragraph(
            f"Custom Branding: {'Yes' if live_form.custom_branding_enabled else 'No'}"
        )
        doc.add_paragraph(f"Total Entries: {entries.count()}")
        doc.add_paragraph("")

        # ── Size Summary ─────────────────────────────────────────────────
        doc.add_heading("Summary by Size", level=2)
        size_summary = (
            entries.values("size").annotate(total=Count("id")).order_by("size")
        )

        summary_table = doc.add_table(rows=1, cols=2)
        summary_table.style = "Light Grid Accent 1"
        hdr = summary_table.rows[0].cells
        hdr[0].text = "Size"
        hdr[1].text = "Total"

        for row_data in size_summary:
            row = summary_table.add_row().cells
            row[0].text = row_data["size"]
            row[1].text = str(row_data["total"])

        doc.add_paragraph("")

        # ── Custom Names by Size (only when custom_branding_enabled) ────
        if live_form.custom_branding_enabled:
            doc.add_heading("Custom Names by Size", level=2)

            # Collect custom names grouped by size
            size_order = ["S", "M", "L", "XL", "XXL", "XXXL", "XXXXL"]
            custom_by_size = {}
            for entry in entries:
                size = entry.size
                if size not in custom_by_size:
                    custom_by_size[size] = []
                custom_by_size[size].append(entry.custom_name or "—")

            custom_table = doc.add_table(rows=1, cols=2)
            custom_table.style = "Light Grid Accent 1"
            hdr_cells = custom_table.rows[0].cells
            hdr_cells[0].text = "Size"
            hdr_cells[1].text = "Custom Names"

            for size in size_order:
                if size in custom_by_size:
                    row = custom_table.add_row().cells
                    row[0].text = size
                    row[1].text = ", ".join(custom_by_size[size])

            doc.add_paragraph("")

        # ── Entries Table ────────────────────────────────────────────────
        doc.add_heading("All Entries", level=2)

        # Build dynamic columns based on custom_branding_enabled
        col_headers = ["#", "Full Name", "Size", "Submitted At"]
        if live_form.custom_branding_enabled:
            col_headers.insert(2, "Custom Name")  # after Full Name

        entry_table = doc.add_table(rows=1, cols=len(col_headers))
        entry_table.style = "Light Grid Accent 1"
        hdr_cells = entry_table.rows[0].cells
        for i, header in enumerate(col_headers):
            hdr_cells[i].text = header

        for entry in entries:
            row_cells = entry_table.add_row().cells
            if live_form.custom_branding_enabled:
                row_cells[0].text = str(entry.serial_number)
                row_cells[1].text = entry.full_name
                row_cells[2].text = entry.custom_name
                row_cells[3].text = entry.size
                row_cells[4].text = entry.created_at.strftime("%Y-%m-%d %H:%M")
            else:
                row_cells[0].text = str(entry.serial_number)
                row_cells[1].text = entry.full_name
                row_cells[2].text = entry.size
                row_cells[3].text = entry.created_at.strftime("%Y-%m-%d %H:%M")

        # ── Save ─────────────────────────────────────────────────────────
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        response = HttpResponse(
            buffer.getvalue(),
            content_type=(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ),
        )
        filename = (
            f"live_form_{live_form.slug}_{timezone.now().strftime('%Y%m%d')}.docx"
        )
        response["Content-Disposition"] = f'attachment; filename="{filename}"'

        logger.info(f"Generated Word document for live form: {live_form.slug}")
        return response

    except Exception as e:
        logger.error(
            f"Error generating Word document for live form {live_form}: {str(e)}"
        )
        raise


# ---------------------------------------------------------------------------
# Excel  (≡ generate_bulk_order_excel)
# ---------------------------------------------------------------------------

def generate_live_form_excel(live_form):
    """
    Generate Excel spreadsheet (.xlsx) for a live form.
    Used by both admin interface and API endpoints.

    custom_name column included only when custom_branding_enabled=True.

    Args:
        live_form: LiveFormLink instance or slug string

    Returns:
        HttpResponse with Excel content

    Raises:
        Exception: For generation errors
    """
    try:
        import xlsxwriter
        from io import BytesIO

        live_form = _get_live_form_with_entries(live_form)
        entries = live_form.entries.all().order_by("serial_number")

        buffer = BytesIO()
        workbook = xlsxwriter.Workbook(buffer, {"in_memory": True})

        # ── Formats ──────────────────────────────────────────────────────
        header_fmt = workbook.add_format(
            {
                "bold": True,
                "bg_color": "#064E3B",  # project primary green
                "font_color": "#FFFFFF",
                "align": "center",
                "valign": "vcenter",
                "border": 1,
            }
        )
        cell_fmt = workbook.add_format(
            {"border": 1, "align": "left", "valign": "vcenter"}
        )
        alt_cell_fmt = workbook.add_format(
            {
                "border": 1,
                "align": "left",
                "valign": "vcenter",
                "bg_color": "#F0FDF4",
            }
        )
        title_fmt = workbook.add_format(
            {
                "bold": True,
                "font_size": 14,
                "font_color": "#064E3B",
            }
        )
        meta_fmt = workbook.add_format({"italic": True, "font_color": "#6B7280"})
        summary_header_fmt = workbook.add_format(
            {
                "bold": True,
                "bg_color": "#F59E0B",  # project accent amber
                "font_color": "#1F2937",
                "border": 1,
                "align": "center",
            }
        )

        # ── Main Entries Sheet ────────────────────────────────────────────
        ws = workbook.add_worksheet("Entries")
        ws.set_zoom(90)

        # Title block
        ws.write(0, 0, settings.COMPANY_NAME, title_fmt)
        ws.write(1, 0, f"Live Form: {live_form.organization_name}", title_fmt)
        ws.write(
            2, 0, f"Generated: {timezone.now().strftime('%B %d, %Y %I:%M %p')}", meta_fmt
        )
        ws.write(
            3, 0, f"Expires: {live_form.expires_at.strftime('%B %d, %Y %I:%M %p')}", meta_fmt
        )
        ws.write(4, 0, f"Total Entries: {entries.count()}", meta_fmt)

        # Column headers (row 6, 0-indexed = row 7 in Excel)
        DATA_ROW_START = 6
        if live_form.custom_branding_enabled:
            headers = ["#", "Full Name", "Custom Name", "Size", "Submitted At"]
            col_widths = [6, 28, 28, 10, 22]
        else:
            headers = ["#", "Full Name", "Size", "Submitted At"]
            col_widths = [6, 30, 10, 22]

        for col_idx, (header, width) in enumerate(zip(headers, col_widths)):
            ws.write(DATA_ROW_START, col_idx, header, header_fmt)
            ws.set_column(col_idx, col_idx, width)

        ws.set_row(DATA_ROW_START, 20)

        # Data rows
        for row_offset, entry in enumerate(entries):
            excel_row = DATA_ROW_START + 1 + row_offset
            fmt = cell_fmt if row_offset % 2 == 0 else alt_cell_fmt

            if live_form.custom_branding_enabled:
                ws.write(excel_row, 0, entry.serial_number, fmt)
                ws.write(excel_row, 1, entry.full_name, fmt)
                ws.write(excel_row, 2, entry.custom_name, fmt)
                ws.write(excel_row, 3, entry.size, fmt)
                ws.write(
                    excel_row,
                    4,
                    entry.created_at.strftime("%Y-%m-%d %H:%M"),
                    fmt,
                )
            else:
                ws.write(excel_row, 0, entry.serial_number, fmt)
                ws.write(excel_row, 1, entry.full_name, fmt)
                ws.write(excel_row, 2, entry.size, fmt)
                ws.write(
                    excel_row,
                    3,
                    entry.created_at.strftime("%Y-%m-%d %H:%M"),
                    fmt,
                )

        ws.freeze_panes(DATA_ROW_START + 1, 0)

        # ── Size Summary Sheet ────────────────────────────────────────────
        ws_summary = workbook.add_worksheet("Size Summary")
        ws_summary.set_zoom(90)
        ws_summary.set_column(0, 0, 14)
        ws_summary.set_column(1, 1, 14)

        ws_summary.write(0, 0, "Size", summary_header_fmt)
        ws_summary.write(0, 1, "Total", summary_header_fmt)

        size_summary = (
            entries.values("size").annotate(count=Count("size")).order_by("size")
        )

        for i, row_data in enumerate(size_summary, start=1):
            fmt = cell_fmt if i % 2 == 0 else alt_cell_fmt
            ws_summary.write(i, 0, row_data["size"], fmt)
            ws_summary.write(i, 1, row_data["count"], fmt)

        # ── Custom Names by Size Sheet (only when custom_branding_enabled) ─
        if live_form.custom_branding_enabled:
            ws_custom = workbook.add_worksheet("Custom Names by Size")
            ws_custom.set_zoom(90)

            # Sheet title block
            ws_custom.write(0, 0, settings.COMPANY_NAME, title_fmt)
            ws_custom.write(
                1, 0,
                f"Live Form: {live_form.organization_name} — Custom Names by Size",
                title_fmt,
            )
            ws_custom.write(
                2, 0,
                f"Generated: {timezone.now().strftime('%B %d, %Y %I:%M %p')}",
                meta_fmt,
            )

            # Column headers
            ws_custom.write(4, 0, "Size", summary_header_fmt)
            ws_custom.write(4, 1, "Custom Name", summary_header_fmt)
            ws_custom.write(4, 2, "Full Name", summary_header_fmt)
            ws_custom.set_column(0, 0, 10)
            ws_custom.set_column(1, 1, 32)
            ws_custom.set_column(2, 2, 32)
            ws_custom.set_row(4, 20)

            # Sort entries by size (SIZE_CHOICES order) then full_name
            size_order = ["S", "M", "L", "XL", "XXL", "XXXL", "XXXXL"]
            sorted_entries = sorted(
                entries,
                key=lambda e: (
                    size_order.index(e.size) if e.size in size_order else 99,
                    e.full_name,
                ),
            )

            # Group rows by size with a subtle size divider
            size_divider_fmt = workbook.add_format(
                {
                    "bold": True,
                    "bg_color": "#D1FAE5",
                    "font_color": "#064E3B",
                    "border": 1,
                    "align": "center",
                    "valign": "vcenter",
                }
            )

            current_size = None
            excel_row = 5

            for entry in sorted_entries:
                # Insert a size-group divider row whenever size changes
                if entry.size != current_size:
                    current_size = entry.size
                    ws_custom.write(excel_row, 0, current_size, size_divider_fmt)
                    ws_custom.write(excel_row, 1, "", size_divider_fmt)
                    ws_custom.write(excel_row, 2, "", size_divider_fmt)
                    ws_custom.set_row(excel_row, 18)
                    excel_row += 1

                row_fmt = cell_fmt if excel_row % 2 == 0 else alt_cell_fmt
                ws_custom.write(excel_row, 0, entry.size, row_fmt)
                ws_custom.write(excel_row, 1, entry.custom_name or "—", row_fmt)
                ws_custom.write(excel_row, 2, entry.full_name, row_fmt)
                excel_row += 1

            ws_custom.freeze_panes(5, 0)

        workbook.close()
        buffer.seek(0)

        response = HttpResponse(
            buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
        filename = (
            f"live_form_{live_form.slug}_{timezone.now().strftime('%Y%m%d')}.xlsx"
        )
        response["Content-Disposition"] = f'attachment; filename="{filename}"'

        logger.info(f"Generated Excel for live form: {live_form.slug}")
        return response

    except Exception as e:
        logger.error(
            f"Error generating Excel for live form {live_form}: {str(e)}"
        )
        raise
