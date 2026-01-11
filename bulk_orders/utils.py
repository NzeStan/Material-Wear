# bulk_orders/utils.py
"""
Centralized utility functions for bulk orders app.
Contains document generation (PDF, Word, Excel) and coupon management.
"""
import string
import random
import logging
from django.utils import timezone
from django.template.loader import render_to_string
from django.conf import settings
from django.http import HttpResponse
from django.db.models import Count, Q, Prefetch
from django.core.paginator import Paginator
from io import BytesIO
from .models import CouponCode, BulkOrderLink, OrderEntry

logger = logging.getLogger(__name__)


def generate_coupon_codes(bulk_order, count=10):
    """
    Generate unique coupon codes for a bulk order.
    
    Args:
        bulk_order: BulkOrderLink instance
        count: Number of coupons to generate (default 10)
    
    Returns:
        List of CouponCode instances
    
    Raises:
        Exception: If coupon generation fails
    """
    chars = string.ascii_uppercase + string.digits
    codes = []
    try:
        for _ in range(count):
            while True:
                code = "".join(random.choices(chars, k=8))
                if not CouponCode.objects.filter(code=code).exists():
                    coupon = CouponCode.objects.create(bulk_order=bulk_order, code=code)
                    codes.append(coupon)
                    break
        logger.info(f"Generated {count} coupon codes for bulk order: {bulk_order.id}")
        return codes
    except Exception as e:
        logger.error(f"Error generating coupon codes: {str(e)}")
        raise


def _get_bulk_order_with_orders(bulk_order):
    """
    Helper to get bulk order with optimized prefetch.
    Handles both slug string and BulkOrderLink instance.
    
    Args:
        bulk_order: Either BulkOrderLink instance or slug string
    
    Returns:
        BulkOrderLink instance with prefetched orders
    """
    if isinstance(bulk_order, str):
        # Fetch by slug
        return BulkOrderLink.objects.prefetch_related(
            Prefetch(
                "orders",
                queryset=OrderEntry.objects.select_related("coupon_used")
                .filter(Q(paid=True) | Q(coupon_used__isnull=False))
                .order_by("size", "full_name"),
            )
        ).get(slug=bulk_order)
    else:
        # Refetch with proper prefetch
        return BulkOrderLink.objects.prefetch_related(
            Prefetch(
                "orders",
                queryset=OrderEntry.objects.select_related("coupon_used")
                .filter(Q(paid=True) | Q(coupon_used__isnull=False))
                .order_by("size", "full_name"),
            )
        ).get(id=bulk_order.id)


def generate_bulk_order_pdf(bulk_order, request=None):
    """
    Generate PDF summary for a bulk order.
    Used by both admin interface and API endpoints.
    
    Args:
        bulk_order: BulkOrderLink instance or slug string
        request: Optional request object for building absolute URIs
    
    Returns:
        HttpResponse with PDF content
    
    Raises:
        ImportError: If WeasyPrint is not installed
        Exception: For other PDF generation errors
    """
    try:
        from weasyprint import HTML
        
        bulk_order = _get_bulk_order_with_orders(bulk_order)
        orders = bulk_order.orders.all()
        size_summary = orders.values("size").annotate(count=Count("size")).order_by("size")
        
        context = {
            'bulk_order': bulk_order,
            'size_summary': size_summary,
            'orders': orders,
            'total_orders': orders.count(),
            'company_name': settings.COMPANY_NAME,
            'company_address': settings.COMPANY_ADDRESS,
            'company_phone': settings.COMPANY_PHONE,
            'company_email': settings.COMPANY_EMAIL,
            'now': timezone.now(),
        }
        
        html_string = render_to_string('bulk_orders/pdf_template.html', context)
        
        if request:
            html = HTML(string=html_string, base_url=request.build_absolute_uri())
        else:
            html = HTML(string=html_string)
        
        pdf = html.write_pdf()
        
        response = HttpResponse(pdf, content_type='application/pdf')
        filename = f'bulk_order_{bulk_order.slug}_{timezone.now().strftime("%Y%m%d")}.pdf'
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        logger.info(f"Generated PDF for bulk order: {bulk_order.slug}")
        return response
        
    except ImportError as e:
        logger.error(f"WeasyPrint not available: {str(e)}")
        raise ImportError("PDF generation not available. Install GTK+ libraries for WeasyPrint.")
    except Exception as e:
        logger.error(f"Error generating PDF for {bulk_order}: {str(e)}")
        raise


def generate_bulk_order_word(bulk_order):
    """
    Generate Word document for a bulk order.
    Used by both admin interface and API endpoints.
    
    Args:
        bulk_order: BulkOrderLink instance or slug string
    
    Returns:
        HttpResponse with Word document content
    
    Raises:
        Exception: For document generation errors
    """
    try:
        from docx import Document
        
        bulk_order = _get_bulk_order_with_orders(bulk_order)
        doc = Document()
        
        # Header
        doc.add_heading(settings.COMPANY_NAME, 0)
        doc.add_heading(f'Bulk Order: {bulk_order.organization_name}', level=1)
        doc.add_paragraph(f"Generated: {timezone.now().strftime('%B %d, %Y - %I:%M %p')}")
        doc.add_paragraph(f'Payment Deadline: {bulk_order.payment_deadline.strftime("%B %d, %Y")}')
        doc.add_paragraph(f'Custom Branding: {"Yes" if bulk_order.custom_branding_enabled else "No"}')
        doc.add_paragraph('')
        
        orders = bulk_order.orders.all()
        
        # Size Summary
        doc.add_heading('Summary by Size', level=2)
        size_summary = orders.values("size").annotate(total=Count("id")).order_by("size")
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Size'
        header_cells[1].text = 'Total'
        
        for size_info in size_summary:
            row_cells = table.add_row().cells
            row_cells[0].text = size_info['size']
            row_cells[1].text = str(size_info['total'])
        
        doc.add_paragraph()
        
        # Orders by Size (paginated for large datasets)
        paginator = Paginator(orders, 1000)
        
        for page_num in paginator.page_range:
            page = paginator.page(page_num)
            
            # Group by size
            size_groups = {}
            for order in page.object_list:
                if order.size not in size_groups:
                    size_groups[order.size] = []
                size_groups[order.size].append(order)
            
            # Create table for each size
            for size, size_orders in sorted(size_groups.items()):
                doc.add_heading(f'Size: {size}', level=3)
                
                if bulk_order.custom_branding_enabled:
                    # WITH custom branding
                    table = doc.add_table(rows=1, cols=4)
                    table.style = 'Table Grid'
                    header_cells = table.rows[0].cells
                    header_cells[0].text = 'S/N'
                    header_cells[1].text = 'Name'
                    header_cells[2].text = 'Custom Name'
                    header_cells[3].text = 'Status'
                    
                    for idx, order in enumerate(size_orders, 1):
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(idx)
                        row_cells[1].text = order.full_name
                        row_cells[2].text = order.custom_name or ''
                        row_cells[3].text = 'Coupon' if order.coupon_used else 'Paid'
                else:
                    # WITHOUT custom branding
                    table = doc.add_table(rows=1, cols=3)
                    table.style = 'Table Grid'
                    header_cells = table.rows[0].cells
                    header_cells[0].text = 'S/N'
                    header_cells[1].text = 'Name'
                    header_cells[2].text = 'Status'
                    
                    for idx, order in enumerate(size_orders, 1):
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(idx)
                        row_cells[1].text = order.full_name
                        row_cells[2].text = 'Coupon' if order.coupon_used else 'Paid'
                
                doc.add_paragraph()
        
        # Footer
        doc.add_paragraph("")
        footer_para = doc.add_paragraph()
        footer_para.add_run(f"{settings.COMPANY_NAME}\n").bold = True
        footer_para.add_run(f"{settings.COMPANY_ADDRESS}\n")
        footer_para.add_run(f"ðŸ“ž {settings.COMPANY_PHONE} | ðŸ“§ {settings.COMPANY_EMAIL}")
        
        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        response = HttpResponse(
            buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        filename = f'bulk_order_{bulk_order.slug}_{timezone.now().strftime("%Y%m%d")}.docx'
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        logger.info(f"Generated Word document for bulk order: {bulk_order.slug}")
        return response
        
    except Exception as e:
        logger.error(f"Error generating Word document for {bulk_order}: {str(e)}")
        raise


def generate_bulk_order_excel(bulk_order):
    """
    Generate Excel spreadsheet for a bulk order.
    Used by both admin interface and API endpoints.
    
    Args:
        bulk_order: BulkOrderLink instance or slug string
    
    Returns:
        HttpResponse with Excel content
    
    Raises:
        Exception: For Excel generation errors
    """
    try:
        import xlsxwriter
        
        bulk_order = _get_bulk_order_with_orders(bulk_order)
        output = BytesIO()
        workbook = xlsxwriter.Workbook(output, {'constant_memory': True})
        
        # ====== CELL FORMATS ======
        title_format = workbook.add_format({'bold': True, 'font_size': 16, 'align': 'left'})
        subtitle_format = workbook.add_format({'bold': True, 'font_size': 12, 'align': 'left'})
        info_format = workbook.add_format({'font_size': 10, 'align': 'left'})
        section_header_format = workbook.add_format({
            'bold': True, 'font_size': 12, 'bg_color': '#4472C4',
            'font_color': 'white', 'align': 'center', 'valign': 'vcenter', 'border': 1
        })
        table_header_format = workbook.add_format({
            'bold': True, 'bg_color': '#D9E1F2', 'border': 1,
            'align': 'center', 'valign': 'vcenter'
        })
        cell_format = workbook.add_format({'border': 1, 'align': 'center', 'valign': 'vcenter'})
        cell_left_format = workbook.add_format({'border': 1, 'align': 'left', 'valign': 'vcenter'})
        total_format = workbook.add_format({
            'bold': True, 'bg_color': '#FFF2CC', 'border': 1,
            'align': 'center', 'valign': 'vcenter'
        })
        
        # ====== CREATE WORKSHEET ======
        worksheet_name = bulk_order.organization_name[:31]  # Excel 31 char limit
        worksheet = workbook.add_worksheet(worksheet_name)
        
        # ====== TITLE SECTION ======
        row = 0
        worksheet.write(row, 0, settings.COMPANY_NAME, title_format)
        row += 1
        worksheet.write(row, 0, f"Bulk Order: {bulk_order.organization_name}", subtitle_format)
        row += 1
        worksheet.write(row, 0, f"Generated: {timezone.now().strftime('%B %d, %Y - %I:%M %p')}", info_format)
        row += 1
        worksheet.write(row, 0, f"Payment Deadline: {bulk_order.payment_deadline.strftime('%B %d, %Y')}", info_format)
        row += 1
        worksheet.write(row, 0, f"Custom Branding: {'Yes' if bulk_order.custom_branding_enabled else 'No'}", info_format)
        row += 2
        
        # ====== SIZE SUMMARY SECTION ======
        orders = bulk_order.orders.all()
        size_summary = orders.values("size").annotate(total=Count("id")).order_by("size")
        
        worksheet.merge_range(row, 0, row, 1, 'SUMMARY BY SIZE', section_header_format)
        row += 1
        
        for col, header in enumerate(['Size', 'Total']):
            worksheet.write(row, col, header, table_header_format)
        row += 1
        
        grand_total = 0
        for size_data in size_summary:
            worksheet.write(row, 0, size_data['size'], cell_format)
            worksheet.write(row, 1, size_data['total'], cell_format)
            grand_total += size_data['total']
            row += 1
        
        worksheet.write(row, 0, 'TOTAL', total_format)
        worksheet.write(row, 1, grand_total, total_format)
        row += 2
        
        # ====== ORDER DETAILS SECTION ======
        if bulk_order.custom_branding_enabled:
            worksheet.merge_range(row, 0, row, 4, 'ORDER DETAILS', section_header_format)
            headers = ['S/N', 'Size', 'Name', 'Custom Name', 'Status']
        else:
            worksheet.merge_range(row, 0, row, 3, 'ORDER DETAILS', section_header_format)
            headers = ['S/N', 'Size', 'Name', 'Status']
        row += 1
        
        for col, header in enumerate(headers):
            worksheet.write(row, col, header, table_header_format)
        row += 1
        
        # Write order data
        serial_number = 1
        for order in orders:
            col = 0
            worksheet.write(row, col, serial_number, cell_format)
            col += 1
            worksheet.write(row, col, order.size, cell_format)
            col += 1
            worksheet.write(row, col, order.full_name, cell_left_format)
            col += 1
            
            if bulk_order.custom_branding_enabled:
                worksheet.write(row, col, order.custom_name or '', cell_left_format)
                col += 1
            
            status_text = 'Coupon' if order.coupon_used else 'Paid'
            worksheet.write(row, col, status_text, cell_format)
            
            row += 1
            serial_number += 1
        
        # ====== SET COLUMN WIDTHS ======
        worksheet.set_column(0, 0, 6)   # S/N
        worksheet.set_column(1, 1, 8)   # Size
        worksheet.set_column(2, 2, 30)  # Name
        if bulk_order.custom_branding_enabled:
            worksheet.set_column(3, 3, 30)  # Custom Name
            worksheet.set_column(4, 4, 12)  # Status
        else:
            worksheet.set_column(3, 3, 12)  # Status
        
        # ====== FINALIZE ======
        workbook.close()
        output.seek(0)
        
        response = HttpResponse(
            output.read(),
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
        filename = f'bulk_order_{bulk_order.slug}_{timezone.now().strftime("%Y%m%d")}.xlsx'
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        logger.info(f"Generated Excel for bulk order: {bulk_order.slug}")
        return response
        
    except Exception as e:
        logger.error(f"Error generating Excel for {bulk_order}: {str(e)}")
        raise