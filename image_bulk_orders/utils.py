# image_bulk_orders/utils.py
"""
Utility functions for Image Bulk Orders.
IDENTICAL to bulk_orders/utils.py with image downloading added.
"""
import string
import random
import logging
from django.utils import timezone
from django.template.loader import render_to_string
from django.conf import settings
from django.http import HttpResponse
from django.db.models import Count, Q, Prefetch
from io import BytesIO
import zipfile
import os
import tempfile
import requests
from pathlib import Path
from django.core.paginator import Paginator

from .models import ImageCouponCode, ImageBulkOrderLink, ImageOrderEntry

logger = logging.getLogger(__name__)


def generate_coupon_codes_image(bulk_order, count=10):
    """
    Generate unique coupon codes for an image bulk order.
    
    Args:
        bulk_order: ImageBulkOrderLink instance
        count: Number of coupons to generate (default 10)
    
    Returns:
        List of ImageCouponCode instances
    """
    chars = string.ascii_uppercase + string.digits
    codes = []
    try:
        for _ in range(count):
            while True:
                code = "".join(random.choices(chars, k=8))
                if not ImageCouponCode.objects.filter(code=code).exists():
                    coupon = ImageCouponCode.objects.create(bulk_order=bulk_order, code=code)
                    codes.append(coupon)
                    break
        logger.info(f"Generated {count} coupon codes for image bulk order: {bulk_order.id}")
        return codes
    except Exception as e:
        logger.error(f"Error generating coupon codes: {str(e)}")
        raise


def _get_image_bulk_order_with_orders(bulk_order):
    """
    Helper to get bulk order with optimized prefetch.
    
    ✅ CRITICAL: Orders by size first, then full_name (for proper serial numbering)
    """
    if isinstance(bulk_order, str):
        return ImageBulkOrderLink.objects.prefetch_related(
            Prefetch(
                "orders",
                queryset=ImageOrderEntry.objects.select_related("coupon_used")
                .filter(Q(paid=True) | Q(coupon_used__isnull=False))
                .order_by("size", "full_name"),  # ✅ CRITICAL: Order by size first
            )
        ).get(slug=bulk_order)
    else:
        return ImageBulkOrderLink.objects.prefetch_related(
            Prefetch(
                "orders",
                queryset=ImageOrderEntry.objects.select_related("coupon_used")
                .filter(Q(paid=True) | Q(coupon_used__isnull=False))
                .order_by("size", "full_name"),  # ✅ CRITICAL: Order by size first
            )
        ).get(id=bulk_order.id)


def generate_image_bulk_order_pdf(bulk_order, request=None):
    """
    Generate PDF summary for an image bulk order.
    
    ✅ FIXES:
    - Proper serial number ordering (by size, then name)
    - Correct total counts in summary
    """
    try:
        from weasyprint import HTML
        
        bulk_order = _get_image_bulk_order_with_orders(bulk_order)
        orders = bulk_order.orders.all()
        size_summary = orders.values("size").annotate(count=Count("size")).order_by("size")
        
        # Count paid orders (including coupon users)
        paid_orders = orders.filter(Q(paid=True) | Q(coupon_used__isnull=False))
        total_paid = paid_orders.count()
        
        context = {
            'bulk_order': bulk_order,
            'size_summary': size_summary,
            'orders': orders,
            'paid_orders': paid_orders,  # ✅ FIX: Pass paid_orders to template
            'total_orders': orders.count(),
            'total_paid': total_paid,
            'company_name': settings.COMPANY_NAME,
            'company_address': settings.COMPANY_ADDRESS,
            'company_phone': settings.COMPANY_PHONE,
            'company_email': settings.COMPANY_EMAIL,
            'now': timezone.now(),
        }
        
        html_string = render_to_string('image_bulk_orders/pdf_template.html', context)
        
        if request:
            html = HTML(string=html_string, base_url=request.build_absolute_uri())
        else:
            html = HTML(string=html_string)
        
        pdf = html.write_pdf()
        
        response = HttpResponse(pdf, content_type='application/pdf')
        filename = f'image_bulk_order_{bulk_order.slug}_{timezone.now().strftime("%Y%m%d")}.pdf'
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        logger.info(f"Generated PDF for image bulk order: {bulk_order.slug}")
        return response
        
    except ImportError as e:
        logger.error(f"WeasyPrint not available: {str(e)}")
        raise ImportError("PDF generation not available. Install GTK+ libraries for WeasyPrint.")
    except Exception as e:
        logger.error(f"Error generating PDF: {str(e)}")
        raise


def generate_image_bulk_order_word(bulk_order):
    """
    Generate Word document for an image bulk order.
    
    ✅ NEW: Includes "Custom Names by Size" section at the end
    ✅ FIX: Serial numbers restart at 1 for each size section
    """
    try:
        from docx import Document
        
        bulk_order = _get_image_bulk_order_with_orders(bulk_order)
        doc = Document()
        
        # Header
        doc.add_heading(settings.COMPANY_NAME, 0)
        doc.add_heading(f'Image Bulk Order: {bulk_order.organization_name}', level=1)
        doc.add_paragraph(f"Generated: {timezone.now().strftime('%B %d, %Y - %I:%M %p')}")
        doc.add_paragraph(f'Payment Deadline: {bulk_order.payment_deadline.strftime("%B %d, %Y")}')
        doc.add_paragraph(f'Custom Branding: {"Yes" if bulk_order.custom_branding_enabled else "No"}')
        doc.add_paragraph('')
        
        orders = bulk_order.orders.all()
        
        # Size Summary
        doc.add_heading('Summary by Size', level=2)
        size_summary = orders.values("size").annotate(total=Count("id")).order_by("size")
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Size'
        header_cells[1].text = 'Total'
        
        for size_info in size_summary:
            row_cells = table.add_row().cells
            row_cells[0].text = size_info['size']
            row_cells[1].text = str(size_info['total'])
        
        doc.add_paragraph()
        
        # Orders by Size (paginated for large datasets)
        paginator = Paginator(orders, 1000)
        
        for page_num in paginator.page_range:
            page = paginator.page(page_num)
            page_orders = page.object_list
            
            # Group by size
            sizes_in_page = set(order.size for order in page_orders)
            
            for size in sorted(sizes_in_page):
                size_orders = [order for order in page_orders if order.size == size]
                
                if size_orders:
                    doc.add_heading(f'Size: {size} ({len(size_orders)} people)', level=2)
                    
                    # Determine columns based on custom branding
                    if bulk_order.custom_branding_enabled:
                        table = doc.add_table(rows=1, cols=5)
                        header_cells = table.rows[0].cells
                        header_cells[0].text = 'S/N'
                        header_cells[1].text = 'Name'
                        header_cells[2].text = 'Custom Name'
                        header_cells[3].text = 'Image'
                        header_cells[4].text = 'Status'
                        
                        # ✅ FIX: Serial number restarts at 1 for each size
                        for idx, order in enumerate(size_orders, 1):
                            row_cells = table.add_row().cells
                            row_cells[0].text = str(idx)  # ✅ Starts at 1 for each size
                            row_cells[1].text = order.full_name
                            row_cells[2].text = order.custom_name or '-'
                            row_cells[3].text = '✓' if order.image else '-'
                            row_cells[4].text = 'Coupon' if order.coupon_used else 'Paid'
                    else:
                        table = doc.add_table(rows=1, cols=4)
                        header_cells = table.rows[0].cells
                        header_cells[0].text = 'S/N'
                        header_cells[1].text = 'Name'
                        header_cells[2].text = 'Image'
                        header_cells[3].text = 'Status'
                        
                        # ✅ FIX: Serial number restarts at 1 for each size
                        for idx, order in enumerate(size_orders, 1):
                            row_cells = table.add_row().cells
                            row_cells[0].text = str(idx)  # ✅ Starts at 1 for each size
                            row_cells[1].text = order.full_name
                            row_cells[2].text = '✓' if order.image else '-'
                            row_cells[3].text = 'Coupon' if order.coupon_used else 'Paid'
                    
                    table.style = 'Light Grid Accent 1'
                    doc.add_paragraph()
        
        # ====== NEW SECTION: Custom Names by Size (only if custom branding enabled) ======
        if bulk_order.custom_branding_enabled:
            # Add page break to start on new page
            doc.add_page_break()
            
            doc.add_heading('Custom Names by Size', level=1)
            doc.add_paragraph('This section shows all custom names grouped by size for easy copying.')
            doc.add_paragraph()
            
            # Query orders for THIS bulk_order ONLY with custom names
            bulk_order_orders = ImageOrderEntry.objects.filter(
                bulk_order=bulk_order,
                custom_name__isnull=False
            ).exclude(
                custom_name=''
            ).order_by('size', 'full_name')
            
            # Get size summary ONLY for orders with custom names
            custom_names_size_summary = bulk_order_orders.values('size').annotate(
                count=Count('id')
            ).order_by('size')
            
            # Group orders by size
            for size_info in custom_names_size_summary:
                size = size_info['size']
                size_orders = bulk_order_orders.filter(size=size)
                
                # Get custom names (uppercase for consistency)
                custom_names = [
                    order.custom_name.upper()
                    for order in size_orders
                ]
                
                if custom_names:
                    # Add size header
                    doc.add_heading(f'SIZE: {size}', level=2)
                    
                    # Create table with 5 columns for grid layout
                    num_cols = 5
                    num_rows = (len(custom_names) + num_cols - 1) // num_cols  # Ceiling division
                    
                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    table.style = 'Light Grid Accent 1'
                    
                    # Fill table with custom names
                    for idx, custom_name in enumerate(custom_names):
                        row_idx = idx // num_cols
                        col_idx = idx % num_cols
                        table.rows[row_idx].cells[col_idx].text = custom_name
                    
                    doc.add_paragraph()
        
        # Footer
        doc.add_paragraph("")
        footer_para = doc.add_paragraph()
        footer_para.add_run(f"{settings.COMPANY_NAME}\n").bold = True
        footer_para.add_run(f"{settings.COMPANY_ADDRESS}\n")
        footer_para.add_run(f"📞 {settings.COMPANY_PHONE} | 📧 {settings.COMPANY_EMAIL}")
        
        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        response = HttpResponse(
            buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        filename = f'image_bulk_order_{bulk_order.slug}_{timezone.now().strftime("%Y%m%d")}.docx'
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        logger.info(f"Generated Word document for image bulk order: {bulk_order.slug}")
        return response
        
    except Exception as e:
        logger.error(f"Error generating Word document: {str(e)}")
        raise


def generate_image_bulk_order_excel(bulk_order):
    """
    Generate Excel spreadsheet for an image bulk order.
    
    ✅ FIX: Serial numbers increase continuously across all sizes
    """
    try:
        import xlsxwriter
        
        bulk_order = _get_image_bulk_order_with_orders(bulk_order)
        output = BytesIO()
        workbook = xlsxwriter.Workbook(output, {'constant_memory': True})
        
        # ====== CELL FORMATS ======
        title_format = workbook.add_format({'bold': True, 'font_size': 16, 'align': 'left'})
        subtitle_format = workbook.add_format({'bold': True, 'font_size': 12, 'align': 'left'})
        info_format = workbook.add_format({'font_size': 10, 'align': 'left'})
        section_header_format = workbook.add_format({
            'bold': True, 'font_size': 12, 'bg_color': '#4472C4',
            'font_color': 'white', 'align': 'center', 'valign': 'vcenter', 'border': 1
        })
        table_header_format = workbook.add_format({
            'bold': True, 'bg_color': '#D9E1F2', 'border': 1,
            'align': 'center', 'valign': 'vcenter'
        })
        cell_format = workbook.add_format({'border': 1, 'align': 'center', 'valign': 'vcenter'})
        cell_left_format = workbook.add_format({'border': 1, 'align': 'left', 'valign': 'vcenter'})
        total_format = workbook.add_format({
            'bold': True, 'bg_color': '#FFF2CC', 'border': 1,
            'align': 'center', 'valign': 'vcenter'
        })
        
        # ====== CREATE WORKSHEET ======
        worksheet_name = bulk_order.organization_name[:31]  # Excel 31 char limit
        worksheet = workbook.add_worksheet(worksheet_name)
        
        # ====== TITLE SECTION ======
        row = 0
        worksheet.write(row, 0, settings.COMPANY_NAME, title_format)
        row += 1
        worksheet.write(row, 0, f"Image Bulk Order: {bulk_order.organization_name}", subtitle_format)
        row += 2
        worksheet.write(row, 0, f"Generated: {timezone.now().strftime('%B %d, %Y - %I:%M %p')}", info_format)
        row += 1
        worksheet.write(row, 0, f"Deadline: {bulk_order.payment_deadline.strftime('%B %d, %Y')}", info_format)
        row += 2
        
        # ====== SIZE SUMMARY ======
        worksheet.write(row, 0, "SIZE SUMMARY", section_header_format)
        worksheet.write(row, 1, "", section_header_format)
        row += 1
        
        worksheet.write(row, 0, "Size", table_header_format)
        worksheet.write(row, 1, "Total", table_header_format)
        row += 1
        
        orders = bulk_order.orders.all()
        size_summary = orders.values("size").annotate(total=Count("id")).order_by("size")
        
        for size_info in size_summary:
            worksheet.write(row, 0, size_info['size'], cell_format)
            worksheet.write(row, 1, size_info['total'], cell_format)
            row += 1
        
        # Total row
        worksheet.write(row, 0, "TOTAL", total_format)
        worksheet.write(row, 1, orders.count(), total_format)
        row += 3
        
        # ====== ALL ORDERS TABLE ======
        worksheet.write(row, 0, "ALL ORDERS", section_header_format)
        if bulk_order.custom_branding_enabled:
            for i in range(1, 5):
                worksheet.write(row, i, "", section_header_format)
        else:
            for i in range(1, 4):
                worksheet.write(row, i, "", section_header_format)
        row += 1
        
        # Table headers
        col = 0
        worksheet.write(row, col, "S/N", table_header_format)
        col += 1
        worksheet.write(row, col, "Size", table_header_format)
        col += 1
        worksheet.write(row, col, "Full Name", table_header_format)
        col += 1
        
        if bulk_order.custom_branding_enabled:
            worksheet.write(row, col, "Custom Name", table_header_format)
            col += 1
        
        worksheet.write(row, col, "Image", table_header_format)
        col += 1
        worksheet.write(row, col, "Status", table_header_format)
        row += 1
        
        # Write order data
        # ✅ FIX: Serial number increases continuously (1, 2, 3... across all sizes)
        serial_number = 1
        for order in orders:
            col = 0
            worksheet.write(row, col, serial_number, cell_format)  # ✅ Continuous numbering
            col += 1
            worksheet.write(row, col, order.size, cell_format)
            col += 1
            worksheet.write(row, col, order.full_name, cell_left_format)
            col += 1
            
            if bulk_order.custom_branding_enabled:
                worksheet.write(row, col, order.custom_name or '', cell_left_format)
                col += 1
            
            worksheet.write(row, col, '✓' if order.image else '-', cell_format)
            col += 1
            
            status_text = 'Coupon' if order.coupon_used else 'Paid'
            worksheet.write(row, col, status_text, cell_format)
            
            row += 1
            serial_number += 1  # ✅ Increment for next order
        
        # ====== SET COLUMN WIDTHS ======
        worksheet.set_column(0, 0, 6)   # S/N
        worksheet.set_column(1, 1, 8)   # Size
        worksheet.set_column(2, 2, 30)  # Name
        if bulk_order.custom_branding_enabled:
            worksheet.set_column(3, 3, 30)  # Custom Name
            worksheet.set_column(4, 4, 8)   # Image
            worksheet.set_column(5, 5, 12)  # Status
        else:
            worksheet.set_column(3, 3, 8)   # Image
            worksheet.set_column(4, 4, 12)  # Status
        
        # ====== FINALIZE ======
        workbook.close()
        output.seek(0)
        
        response = HttpResponse(
            output.read(),
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
        filename = f'image_bulk_order_{bulk_order.slug}_{timezone.now().strftime("%Y%m%d")}.xlsx'
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        logger.info(f"Generated Excel for image bulk order: {bulk_order.slug}")
        return response
        
    except Exception as e:
        logger.error(f"Error generating Excel: {str(e)}")
        raise


def download_image_from_cloudinary(url):
    """Download image from Cloudinary URL"""
    try:
        response = requests.get(url, timeout=30)
        response.raise_for_status()
        return response.content
    except Exception as e:
        logger.error(f"Error downloading image from {url}: {str(e)}")
        return None


def generate_admin_package_with_images(bulk_order_id):
    """
    Generate complete admin package: PDF + Word + Excel + Images by size.
    
    Creates ZIP structure:
    /package_slug_20240212/
        bulk_order_slug.pdf
        bulk_order_slug.docx
        bulk_order_slug.xlsx
        /images/
            /S/
                001_John_Doe.jpg
                002_Jane_Smith.png
            /M/
                003_Bob_Wilson.jpg
            ...
    """
    try:
        bulk_order = ImageBulkOrderLink.objects.prefetch_related(
            'orders'
        ).get(id=bulk_order_id)
        
        # Create temp directory
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            package_name = f"{bulk_order.slug}_{timezone.now().strftime('%Y%m%d')}"
            package_dir = temp_path / package_name
            package_dir.mkdir()
            
            # Generate documents
            pdf_response = generate_image_bulk_order_pdf(bulk_order)
            word_response = generate_image_bulk_order_word(bulk_order)
            excel_response = generate_image_bulk_order_excel(bulk_order)
            
            # Save documents
            (package_dir / f"{bulk_order.slug}.pdf").write_bytes(pdf_response.content)
            (package_dir / f"{bulk_order.slug}.docx").write_bytes(word_response.content)
            (package_dir / f"{bulk_order.slug}.xlsx").write_bytes(excel_response.content)
            
            # Create images directory structure
            images_dir = package_dir / "images"
            images_dir.mkdir()
            
            # Download and organize images by size
            orders_with_images = bulk_order.orders.filter(image__isnull=False)
            
            for order in orders_with_images:
                # Create size subdirectory
                size_dir = images_dir / order.size
                size_dir.mkdir(exist_ok=True)
                
                # Determine filename
                if order.custom_name:
                    filename_base = order.custom_name.replace(' ', '_')
                else:
                    filename_base = f"{order.serial_number}_{order.full_name.replace(' ', '_')}"
                
                # Get file extension from Cloudinary URL
                image_url = order.image.url
                ext = Path(image_url).suffix or '.jpg'
                filename = f"{filename_base}{ext}"
                
                # Download image
                image_data = download_image_from_cloudinary(image_url)
                if image_data:
                    (size_dir / filename).write_bytes(image_data)
                    logger.info(f"Downloaded image: {filename}")
            
            # Create ZIP
            zip_buffer = BytesIO()
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                for file_path in package_dir.rglob('*'):
                    if file_path.is_file():
                        arcname = file_path.relative_to(package_dir.parent)
                        zip_file.write(file_path, arcname)
            
            zip_buffer.seek(0)
            
            response = HttpResponse(zip_buffer.getvalue(), content_type='application/zip')
            response['Content-Disposition'] = f'attachment; filename="{package_name}.zip"'
            
            logger.info(f"Generated admin package with images for: {bulk_order.slug}")
            return response
            
    except Exception as e:
        logger.error(f"Error generating admin package: {str(e)}")
        raise
