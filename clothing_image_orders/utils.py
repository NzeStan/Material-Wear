# clothing_image_orders/utils.py
"""
Utility functions for clothing image orders.

Includes:
- Coupon code generation
- Payment integration (Paystack)
- Document generation (PDF, Word, Excel)
- Complete package generation with images organized by size
"""
import string
import random
import logging
import requests
import tempfile
import shutil
import os
from pathlib import Path
from io import BytesIO
from decimal import Decimal
from django.utils import timezone
from django.template.loader import render_to_string
from django.conf import settings
from django.http import HttpResponse
from django.db.models import Count, Q
import zipfile

from .models import ClothingImageOrder, ClothingOrderParticipant, ClothingCouponCode

logger = logging.getLogger(__name__)


# ==============================================================================
# COUPON GENERATION
# ==============================================================================

def generate_clothing_coupon_codes(order, count=50):
    """
    Generate unique coupon codes for a clothing order.
    
    Args:
        order: ClothingImageOrder instance
        count: Number of coupons to generate (default 50)
    
    Returns:
        List of ClothingCouponCode instances
    """
    chars = string.ascii_uppercase + string.digits
    codes = []
    
    try:
        for _ in range(count):
            while True:
                code = "".join(random.choices(chars, k=8))
                if not ClothingCouponCode.objects.filter(code=code).exists():
                    coupon = ClothingCouponCode.objects.create(order=order, code=code)
                    codes.append(coupon)
                    break
        
        logger.info(f"Generated {count} coupon codes for order: {order.reference}")
        return codes
        
    except Exception as e:
        logger.error(f"Error generating coupon codes: {str(e)}")
        raise


# ==============================================================================
# PAYMENT INTEGRATION (PAYSTACK)
# ==============================================================================

def initialize_payment(amount, email, reference, callback_url):
    """
    Initialize payment with Paystack.
    
    Args:
        amount: Payment amount (Decimal)
        email: Customer email
        reference: Payment reference
        callback_url: URL to redirect after payment
    
    Returns:
        dict: Paystack API response
    """
    url = "https://api.paystack.co/transaction/initialize"
    
    headers = {
        "Authorization": f"Bearer {settings.PAYSTACK_SECRET_KEY}",
        "Content-Type": "application/json"
    }
    
    # Convert amount to kobo (Paystack uses kobo)
    amount_kobo = int(amount * 100)
    
    payload = {
        "email": email,
        "amount": amount_kobo,
        "reference": reference,
        "callback_url": callback_url,
        "currency": "NGN",
        "metadata": {
            "custom_fields": [
                {
                    "display_name": "Payment For",
                    "variable_name": "payment_for",
                    "value": "Clothing Order"
                }
            ]
        }
    }
    
    try:
        response = requests.post(url, json=payload, headers=headers, timeout=10)
        response.raise_for_status()
        return response.json()
        
    except requests.exceptions.RequestException as e:
        logger.error(f"Paystack initialization error: {str(e)}")
        return None


def verify_payment(reference):
    """
    Verify payment with Paystack.
    
    Args:
        reference: Payment reference
    
    Returns:
        dict: Paystack API response
    """
    url = f"https://api.paystack.co/transaction/verify/{reference}"
    
    headers = {
        "Authorization": f"Bearer {settings.PAYSTACK_SECRET_KEY}",
        "Content-Type": "application/json"
    }
    
    try:
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        return response.json()
        
    except requests.exceptions.RequestException as e:
        logger.error(f"Paystack verification error: {str(e)}")
        return None


# ==============================================================================
# DOCUMENT GENERATION - PDF
# ==============================================================================

def generate_clothing_order_pdf(order):
    """
    Generate PDF summary for a clothing order.
    
    Args:
        order: ClothingImageOrder instance
    
    Returns:
        BytesIO: PDF file buffer
    """
    try:
        from weasyprint import HTML
        
        # Get participants (paid or with coupons)
        participants = order.participants.filter(
            Q(paid=True) | Q(coupon_used__isnull=False)
        ).select_related('coupon_used').order_by('size', 'full_name')
        
        # Size summary
        size_summary = participants.values('size').annotate(
            count=Count('id')
        ).order_by('size')
        
        context = {
            'order': order,
            'participants': participants,
            'size_summary': size_summary,
            'total_participants': participants.count(),
            'company_name': settings.COMPANY_NAME,
            'company_address': settings.COMPANY_ADDRESS,
            'company_phone': settings.COMPANY_PHONE,
            'company_email': settings.COMPANY_EMAIL,
            'generated_date': timezone.now(),
        }
        
        html_string = render_to_string('clothing_image_orders/pdf_template.html', context)
        html = HTML(string=html_string)
        pdf = html.write_pdf()
        
        pdf_buffer = BytesIO(pdf)
        logger.info(f"Generated PDF for order: {order.reference}")
        
        return pdf_buffer
        
    except ImportError:
        logger.error("WeasyPrint not available")
        raise ImportError("PDF generation requires WeasyPrint and GTK+ libraries")
    
    except Exception as e:
        logger.error(f"Error generating PDF: {str(e)}")
        raise


# ==============================================================================
# DOCUMENT GENERATION - WORD
# ==============================================================================

def generate_clothing_order_word(order):
    """
    Generate Word document for a clothing order.
    
    Args:
        order: ClothingImageOrder instance
    
    Returns:
        BytesIO: Word document buffer
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Header
        title = doc.add_heading(settings.COMPANY_NAME, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_heading(f'Clothing Order: {order.organization_name}', level=1)
        doc.add_paragraph(f"Order Reference: {order.reference}")
        doc.add_paragraph(f"Title: {order.title}")
        doc.add_paragraph(f"Generated: {timezone.now().strftime('%B %d, %Y - %I:%M %p')}")
        doc.add_paragraph('')
        
        # Get participants
        participants = order.participants.filter(
            Q(paid=True) | Q(coupon_used__isnull=False)
        ).select_related('coupon_used').order_by('size', 'full_name')
        
        # Size Summary
        doc.add_heading('Summary by Size', level=2)
        size_summary = participants.values('size').annotate(
            count=Count('id')
        ).order_by('size')
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Size'
        header_cells[1].text = 'Count'
        
        for size_info in size_summary:
            row_cells = table.add_row().cells
            row_cells[0].text = size_info['size']
            row_cells[1].text = str(size_info['count'])
        
        doc.add_paragraph('')
        
        # Participants List
        doc.add_heading('Participants', level=2)
        
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        header_cells = table.rows[0].cells
        headers = ['S/N', 'Name', 'Size', 'Custom Name', 'Status']
        for i, header in enumerate(headers):
            header_cells[i].text = header
        
        for idx, participant in enumerate(participants, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[1].text = participant.full_name
            row_cells[2].text = participant.size
            row_cells[3].text = participant.custom_name or '-'
            
            if participant.coupon_used:
                row_cells[4].text = 'COUPON'
            else:
                row_cells[4].text = 'PAID'
        
        # Footer
        doc.add_paragraph('')
        footer_para = doc.add_paragraph()
        footer_para.add_run(f"{settings.COMPANY_NAME}\n").bold = True
        footer_para.add_run(f"{settings.COMPANY_ADDRESS}\n")
        footer_para.add_run(f"📞 {settings.COMPANY_PHONE} | 📧 {settings.COMPANY_EMAIL}")
        
        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        logger.info(f"Generated Word document for order: {order.reference}")
        return buffer
        
    except Exception as e:
        logger.error(f"Error generating Word document: {str(e)}")
        raise


# ==============================================================================
# DOCUMENT GENERATION - EXCEL
# ==============================================================================

def generate_clothing_order_excel(order):
    """
    Generate Excel spreadsheet for a clothing order.
    
    Args:
        order: ClothingImageOrder instance
    
    Returns:
        BytesIO: Excel file buffer
    """
    try:
        import xlsxwriter
        
        buffer = BytesIO()
        workbook = xlsxwriter.Workbook(buffer, {'in_memory': True})
        
        # Define formats
        header_format = workbook.add_format({
            'bold': True,
            'bg_color': '#064E3B',
            'font_color': 'white',
            'align': 'center',
            'valign': 'vcenter',
            'border': 1
        })
        
        cell_format = workbook.add_format({
            'align': 'left',
            'valign': 'vcenter',
            'border': 1
        })
        
        coupon_format = workbook.add_format({
            'bg_color': '#F59E0B',
            'bold': True,
            'border': 1
        })
        
        # Get participants
        participants = order.participants.filter(
            Q(paid=True) | Q(coupon_used__isnull=False)
        ).select_related('coupon_used').order_by('size', 'full_name')
        
        # Sheet 1: Participants
        worksheet = workbook.add_worksheet('Participants')
        
        # Headers
        headers = ['S/N', 'Reference', 'Name', 'Email', 'Phone', 'Size', 'Custom Name', 'Status', 'Image']
        for col, header in enumerate(headers):
            worksheet.write(0, col, header, header_format)
        
        # Data
        for row, participant in enumerate(participants, 1):
            worksheet.write(row, 0, row, cell_format)
            worksheet.write(row, 1, participant.reference, cell_format)
            worksheet.write(row, 2, participant.full_name, cell_format)
            worksheet.write(row, 3, participant.email, cell_format)
            worksheet.write(row, 4, participant.phone or '-', cell_format)
            worksheet.write(row, 5, participant.size, cell_format)
            worksheet.write(row, 6, participant.custom_name or '-', cell_format)
            
            if participant.coupon_used:
                worksheet.write(row, 7, 'COUPON', coupon_format)
            else:
                worksheet.write(row, 7, 'PAID', cell_format)
            
            worksheet.write(row, 8, 'Yes' if participant.image else 'No', cell_format)
        
        # Adjust column widths
        worksheet.set_column('A:A', 6)
        worksheet.set_column('B:B', 18)
        worksheet.set_column('C:C', 25)
        worksheet.set_column('D:D', 30)
        worksheet.set_column('E:E', 15)
        worksheet.set_column('F:F', 8)
        worksheet.set_column('G:G', 20)
        worksheet.set_column('H:H', 10)
        worksheet.set_column('I:I', 8)
        
        # Sheet 2: Size Summary
        summary_sheet = workbook.add_worksheet('Size Summary')
        
        summary_sheet.write('A1', 'Size', header_format)
        summary_sheet.write('B1', 'Count', header_format)
        
        size_summary = participants.values('size').annotate(
            count=Count('id')
        ).order_by('size')
        
        for row, size_info in enumerate(size_summary, 1):
            summary_sheet.write(row, 0, size_info['size'], cell_format)
            summary_sheet.write(row, 1, size_info['count'], cell_format)
        
        summary_sheet.set_column('A:A', 12)
        summary_sheet.set_column('B:B', 12)
        
        workbook.close()
        buffer.seek(0)
        
        logger.info(f"Generated Excel for order: {order.reference}")
        return buffer
        
    except Exception as e:
        logger.error(f"Error generating Excel: {str(e)}")
        raise


# ==============================================================================
# COMPLETE PACKAGE GENERATION WITH IMAGES
# ==============================================================================

def download_cloudinary_image(cloudinary_url, save_path):
    """
    Download image from Cloudinary URL to local path.
    
    Args:
        cloudinary_url: URL of image on Cloudinary
        save_path: Local path to save the image
    
    Returns:
        bool: True if successful, False otherwise
    """
    try:
        response = requests.get(cloudinary_url, timeout=30, stream=True)
        response.raise_for_status()
        
        with open(save_path, 'wb') as f:
            for chunk in response.iter_content(chunk_size=8192):
                f.write(chunk)
        
        return True
        
    except Exception as e:
        logger.error(f"Error downloading image from {cloudinary_url}: {str(e)}")
        return False


def generate_complete_package(order):
    """
    Generate complete document package with organized images.
    
    Creates a directory structure:
    order_{reference}_{org_name}/
    ├── order_details.pdf
    ├── order_details.docx
    ├── order_details.xlsx
    └── images/
        ├── size_S/
        │   ├── 001_JOHN_DOE.jpg
        │   ├── PASTOR_JOHN.jpg
        ├── size_M/
        │   ├── 002_JANE_SMITH.jpg
        ├── size_L/
        └── size_XL/
    
    Args:
        order: ClothingImageOrder instance
    
    Returns:
        BytesIO: ZIP file containing the complete package
    """
    try:
        # Create temporary directory
        temp_dir = tempfile.mkdtemp(prefix='clothing_order_')
        
        # Create main directory
        org_name_safe = order.organization_name.replace(' ', '_')
        package_dir = os.path.join(temp_dir, f"{order.reference}_{org_name_safe}")
        os.makedirs(package_dir, exist_ok=True)
        
        # Generate documents
        logger.info(f"Generating documents for {order.reference}")
        
        # PDF
        try:
            pdf_buffer = generate_clothing_order_pdf(order)
            pdf_path = os.path.join(package_dir, 'order_details.pdf')
            with open(pdf_path, 'wb') as f:
                f.write(pdf_buffer.getvalue())
        except Exception as e:
            logger.error(f"PDF generation failed: {str(e)}")
        
        # Word
        try:
            word_buffer = generate_clothing_order_word(order)
            word_path = os.path.join(package_dir, 'order_details.docx')
            with open(word_path, 'wb') as f:
                f.write(word_buffer.getvalue())
        except Exception as e:
            logger.error(f"Word generation failed: {str(e)}")
        
        # Excel
        try:
            excel_buffer = generate_clothing_order_excel(order)
            excel_path = os.path.join(package_dir, 'order_details.xlsx')
            with open(excel_path, 'wb') as f:
                f.write(excel_buffer.getvalue())
        except Exception as e:
            logger.error(f"Excel generation failed: {str(e)}")
        
        # Create images directory
        images_dir = os.path.join(package_dir, 'images')
        os.makedirs(images_dir, exist_ok=True)
        
        # Get participants with images
        participants = order.participants.filter(
            Q(paid=True) | Q(coupon_used__isnull=False)
        ).exclude(image='').order_by('size', 'serial_number')
        
        logger.info(f"Processing {participants.count()} participant images")
        
        # Download and organize images by size
        for participant in participants:
            # Create size directory
            size_dir = os.path.join(images_dir, f"size_{participant.size}")
            os.makedirs(size_dir, exist_ok=True)
            
            # Get image filename
            filename = participant.get_image_filename()
            
            # Download image
            if participant.image:
                image_url = participant.image.url
                save_path = os.path.join(size_dir, filename)
                
                success = download_cloudinary_image(image_url, save_path)
                if success:
                    logger.debug(f"Downloaded image for {participant.reference}")
                else:
                    logger.warning(f"Failed to download image for {participant.reference}")
        
        # Create ZIP file
        logger.info(f"Creating ZIP package for {order.reference}")
        
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            # Walk through directory and add all files
            for root, dirs, files in os.walk(package_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, temp_dir)
                    zip_file.write(file_path, arcname)
        
        zip_buffer.seek(0)
        
        # Cleanup
        shutil.rmtree(temp_dir)
        
        logger.info(f"Successfully generated complete package for {order.reference}")
        return zip_buffer
        
    except Exception as e:
        logger.error(f"Error generating complete package: {str(e)}")
        # Cleanup on error
        if 'temp_dir' in locals():
            shutil.rmtree(temp_dir, ignore_errors=True)
        raise